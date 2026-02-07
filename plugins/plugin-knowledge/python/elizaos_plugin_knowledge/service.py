from __future__ import annotations

import base64
import hashlib
import logging
import math
import os
import re
import time
import uuid
from typing import Protocol

from elizaos_plugin_knowledge.types import (
    AddKnowledgeOptions,
    ChunkResult,
    CONTEXTUAL_CHUNK_ENRICHMENT_PROMPT_TEMPLATE,
    CONTEXT_TARGETS,
    ContextTargets,
    DEFAULT_CHARS_PER_TOKEN,
    EmbeddingResult,
    KnowledgeConfig,
    KnowledgeDocument,
    KnowledgeFragment,
    KnowledgeItem,
    LoadResult,
    MemoryType,
    PendingRAGEntry,
    ProcessingResult,
    RAGMetadata,
    RetrievedFragmentInfo,
    SearchResult,
)

logger = logging.getLogger(__name__)

DEFAULT_CHUNK_SIZE = 500
DEFAULT_CHUNK_OVERLAP = 100
CHARS_PER_TOKEN = 3.5


# ---------------------------------------------------------------------------
# Protocols for pluggable embedding / memory providers
# ---------------------------------------------------------------------------


class EmbeddingProvider(Protocol):
    async def generate_embedding(self, text: str) -> EmbeddingResult: ...

    async def generate_embeddings_batch(self, texts: list[str]) -> list[EmbeddingResult]: ...


class MemoryStore(Protocol):
    async def create_memory(self, memory: dict) -> str: ...

    async def get_memory_by_id(self, memory_id: str) -> dict | None: ...

    async def update_memory(self, memory: dict) -> None: ...

    async def search_memories(
        self,
        embedding: list[float],
        count: int = 10,
        threshold: float = 0.1,
    ) -> list[dict]: ...

    async def get_memories(
        self,
        table_name: str,
        count: int = 10,
    ) -> list[dict]: ...

    async def delete_memory(self, memory_id: str) -> None: ...


class TextGenerationProvider(Protocol):
    """Optional provider for contextual knowledge enrichment (CTX_KNOWLEDGE)."""

    async def generate_text(self, prompt: str, system: str | None = None) -> str: ...


# ---------------------------------------------------------------------------
# KnowledgeService
# ---------------------------------------------------------------------------


class KnowledgeService:
    """Full-featured knowledge service with parity to the TypeScript implementation.

    Supports:
    - Document deduplication via content hashing (SHA-256 + UUID v5)
    - Fragment creation with configurable chunk sizes
    - Embedding-based similarity search
    - RAG enrichment (format context with metadata for LLM consumption)
    - PDF text extraction
    - Contextual knowledge flag (CTX_KNOWLEDGE_ENABLED equivalent)
    """

    def __init__(
        self,
        config: KnowledgeConfig | None = None,
        embedding_provider: EmbeddingProvider | None = None,
        memory_store: MemoryStore | None = None,
        text_generation_provider: TextGenerationProvider | None = None,
    ) -> None:
        self.config = config or KnowledgeConfig()
        self._embedding_provider = embedding_provider
        self._memory_store = memory_store
        self._text_generation_provider = text_generation_provider
        self._documents: dict[str, KnowledgeDocument] = {}
        self._fragments: dict[str, KnowledgeFragment] = {}
        self._content_hashes: dict[str, str] = {}  # hash -> document_id mapping

        # Pending RAG enrichment queue (mirrors TypeScript pendingRAGEnrichment)
        self._pending_rag_enrichment: list[PendingRAGEntry] = []

        self._load_env_config()

    # ------------------------------------------------------------------
    # Configuration
    # ------------------------------------------------------------------

    def _load_env_config(self) -> None:
        env_mappings: dict[str, str | tuple[str, type | object]] = {
            "EMBEDDING_PROVIDER": "embedding_provider",
            "TEXT_EMBEDDING_MODEL": "embedding_model",
            "EMBEDDING_DIMENSION": ("embedding_dimension", int),
            "OPENAI_API_KEY": "openai_api_key",
            "ANTHROPIC_API_KEY": "anthropic_api_key",
            "GOOGLE_API_KEY": "google_api_key",
            "OPENROUTER_API_KEY": "openrouter_api_key",
            "CTX_KNOWLEDGE_ENABLED": ("ctx_knowledge_enabled", lambda x: x.lower() == "true"),
            "TEXT_PROVIDER": "text_provider",
            "TEXT_MODEL": "text_model",
            "MAX_INPUT_TOKENS": ("max_input_tokens", int),
            "MAX_OUTPUT_TOKENS": ("max_output_tokens", int),
            "KNOWLEDGE_PATH": "knowledge_path",
            "LOAD_DOCS_ON_STARTUP": ("load_docs_on_startup", lambda x: x.lower() != "false"),
        }

        for env_key, attr in env_mappings.items():
            value = os.environ.get(env_key)
            if value is not None:
                if isinstance(attr, tuple):
                    attr_name, converter = attr
                    setattr(self.config, attr_name, converter(value))
                else:
                    setattr(self.config, attr, value)

    # ------------------------------------------------------------------
    # Content-based deduplication
    # ------------------------------------------------------------------

    def _generate_content_id(
        self,
        content: str,
        agent_id: str,
        filename: str | None = None,
        content_type: str | None = None,
        max_chars: int = 2000,
    ) -> str:
        """Generate a deterministic ID based on content hash (SHA-256 + UUID v5).

        Mirrors TypeScript ``generateContentBasedId`` for cross-language parity.
        """
        content_for_hash = content[:max_chars].strip()
        # Normalize line endings
        content_for_hash = content_for_hash.replace("\r\n", "\n").replace("\r", "\n").strip()

        components = [agent_id, content_for_hash]
        if filename:
            components.append(filename)

        hash_input = "::".join(filter(None, components))
        content_hash = hashlib.sha256(hash_input.encode()).hexdigest()

        namespace = uuid.UUID("6ba7b810-9dad-11d1-80b4-00c04fd430c8")
        return str(uuid.uuid5(namespace, content_hash))

    def _compute_content_hash(self, content: str) -> str:
        """Compute a SHA-256 hex digest of the content for fast dedup lookups."""
        normalized = content[:4000].replace("\r\n", "\n").replace("\r", "\n").strip()
        return hashlib.sha256(normalized.encode()).hexdigest()

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def _split_into_chunks(
        self,
        text: str,
        chunk_size: int | None = None,
        overlap: int | None = None,
    ) -> ChunkResult:
        """Split text into overlapping chunks with sentence-boundary awareness."""
        chunk_size = chunk_size or self.config.chunk_size
        overlap = overlap or self.config.chunk_overlap

        char_chunk_size = int(chunk_size * CHARS_PER_TOKEN)
        char_overlap = int(overlap * CHARS_PER_TOKEN)

        chunks: list[str] = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = min(start + char_chunk_size, text_len)

            # Try to break at sentence boundaries
            if end < text_len:
                search_start = start + int(char_chunk_size * 0.8)
                for i in range(end, max(search_start, start), -1):
                    if text[i] in ".!?\n":
                        end = i + 1
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            if end >= text_len:
                break

            new_start = end - char_overlap
            if new_start <= start:
                start = end
            else:
                start = new_start

        return ChunkResult(
            chunks=chunks,
            chunk_count=len(chunks),
            total_tokens=int(len(text) / CHARS_PER_TOKEN),
        )

    # ------------------------------------------------------------------
    # Add knowledge (main entry point)
    # ------------------------------------------------------------------

    async def add_knowledge(
        self,
        options: AddKnowledgeOptions,
    ) -> ProcessingResult:
        """Add a document to the knowledge base with deduplication.

        Mirrors TypeScript ``addKnowledge`` – generates a content-based ID,
        checks for duplicates, then processes document into fragments with
        embeddings.
        """
        agent_id = options.agent_id or "default"

        # Generate content-based document ID for deduplication
        document_id = self._generate_content_id(
            options.content,
            agent_id,
            options.filename,
            options.content_type,
        )

        # Check duplicate via content hash
        content_hash = self._compute_content_hash(options.content)
        if content_hash in self._content_hashes:
            existing_doc_id = self._content_hashes[content_hash]
            existing = self._documents.get(existing_doc_id)
            if existing is not None:
                logger.info(f'"{options.filename}" already exists (hash match) - skipping')
                return ProcessingResult(
                    document_id=existing_doc_id,
                    fragment_count=len(existing.fragments),
                    success=True,
                )

        # Also check by ID (same content + agent + filename)
        if document_id in self._documents:
            existing = self._documents[document_id]
            logger.info(f'"{options.filename}" already exists (ID match) - skipping')
            return ProcessingResult(
                document_id=document_id,
                fragment_count=len(existing.fragments),
                success=True,
            )

        try:
            # Extract text from content
            text_content = await self._extract_text(
                options.content,
                options.content_type,
                options.filename,
            )

            if not text_content or not text_content.strip():
                return ProcessingResult(
                    document_id=document_id,
                    fragment_count=0,
                    success=False,
                    error="No text content extracted",
                )

            # Build document record
            file_ext = options.filename.rsplit(".", 1)[-1].lower() if "." in options.filename else ""
            title = options.filename.replace(f".{file_ext}", "") if file_ext else options.filename

            document = KnowledgeDocument(
                id=document_id,
                content=text_content,
                filename=options.filename,
                content_type=options.content_type,
                file_size=len(options.content),
                content_hash=content_hash,
                metadata={
                    "type": MemoryType.DOCUMENT.value,
                    "source": "knowledge-service",
                    "title": title,
                    "filename": options.filename,
                    "fileExt": file_ext,
                    "fileType": options.content_type,
                    "fileSize": len(options.content),
                    "timestamp": int(time.time() * 1000),
                    **options.metadata,
                },
            )

            # Split into chunks
            chunk_result = self._split_into_chunks(text_content)

            # Create fragments with embeddings
            fragments: list[KnowledgeFragment] = []
            for i, chunk in enumerate(chunk_result.chunks):
                fragment_id = f"{document_id}-fragment-{i}-{int(time.time() * 1000)}"

                # Optionally enrich chunk with contextual knowledge
                enriched_chunk = chunk
                if self.config.ctx_knowledge_enabled and self._text_generation_provider:
                    enriched_chunk = await self._enrich_chunk_with_context(
                        chunk, text_content, options.content_type
                    )

                fragment = KnowledgeFragment(
                    id=fragment_id,
                    document_id=document_id,
                    content=enriched_chunk,
                    position=i,
                    metadata={
                        "type": MemoryType.FRAGMENT.value,
                        "document_id": document_id,
                        "position": i,
                        "timestamp": int(time.time() * 1000),
                        "source": "rag-service-fragment-sync",
                    },
                )
                fragments.append(fragment)
                self._fragments[fragment_id] = fragment

            document.fragments = fragments
            self._documents[document_id] = document
            self._content_hashes[content_hash] = document_id

            # Generate embeddings for all fragments
            if self._embedding_provider:
                await self._generate_fragment_embeddings(fragments)

            # Persist to memory store if available
            if self._memory_store:
                await self._persist_document_to_store(document)

            logger.info(f"Added document '{options.filename}' with {len(fragments)} fragments")

            return ProcessingResult(
                document_id=document_id,
                fragment_count=len(fragments),
                success=True,
            )

        except Exception as e:
            logger.error(f"Error adding knowledge: {e}")
            return ProcessingResult(
                document_id=document_id,
                fragment_count=0,
                success=False,
                error=str(e),
            )

    async def _persist_document_to_store(self, document: KnowledgeDocument) -> None:
        """Persist document and its fragments to the memory store."""
        if not self._memory_store:
            return

        doc_memory = {
            "id": document.id,
            "content": {"text": document.content},
            "metadata": document.metadata,
        }
        await self._memory_store.create_memory(doc_memory)

        for fragment in document.fragments:
            frag_memory: dict[str, object] = {
                "id": fragment.id,
                "content": {"text": fragment.content},
                "metadata": fragment.metadata,
            }
            if fragment.embedding:
                frag_memory["embedding"] = fragment.embedding
            await self._memory_store.create_memory(frag_memory)

    # ------------------------------------------------------------------
    # Text extraction (PDF, DOCX, plain text)
    # ------------------------------------------------------------------

    async def _extract_text(
        self,
        content: str,
        content_type: str,
        filename: str,
    ) -> str:
        """Extract text from various document formats."""
        if content_type.startswith("text/") or content_type in [
            "application/json",
            "application/xml",
        ]:
            # Check if the text content is base64-encoded
            if self._looks_like_base64(content):
                try:
                    decoded = base64.b64decode(content).decode("utf-8")
                    return decoded
                except Exception:
                    pass
            return content

        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            return await self._extract_pdf_text(content, filename)

        if "wordprocessingml" in content_type or filename.lower().endswith(".docx"):
            return await self._extract_docx_text(content)

        # Fallback: try to treat as text
        if self._looks_like_base64(content):
            try:
                return base64.b64decode(content).decode("utf-8")
            except Exception:
                pass

        return content

    async def _extract_pdf_text(self, content: str, filename: str = "") -> str:
        """Extract text from a PDF document.

        Tries PyMuPDF (fitz) first for best extraction quality,
        falls back to pypdf, then to pdfplumber.
        """
        if self._looks_like_base64(content):
            pdf_bytes = base64.b64decode(content)
        else:
            pdf_bytes = content.encode()

        # Strategy 1: PyMuPDF (fitz) - best quality
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            text_parts: list[str] = []
            for page in doc:
                text = page.get_text()
                if text and text.strip():
                    text_parts.append(text.strip())
            doc.close()
            result = "\n\n".join(text_parts)
            if result.strip():
                return self._clean_pdf_text(result)
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"PyMuPDF extraction failed for {filename}: {e}")

        # Strategy 2: pypdf
        try:
            from io import BytesIO

            from pypdf import PdfReader

            reader = PdfReader(BytesIO(pdf_bytes))
            text_parts_pypdf: list[str] = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts_pypdf.append(text)
            result = "\n\n".join(text_parts_pypdf)
            if result.strip():
                return self._clean_pdf_text(result)
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"pypdf extraction failed for {filename}: {e}")

        # Strategy 3: pdfplumber
        try:
            import pdfplumber
            from io import BytesIO

            with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
                text_parts_pb: list[str] = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts_pb.append(text)
            result = "\n\n".join(text_parts_pb)
            if result.strip():
                return self._clean_pdf_text(result)
        except ImportError:
            pass
        except Exception as e:
            logger.debug(f"pdfplumber extraction failed for {filename}: {e}")

        raise ImportError(
            "No PDF extraction library available. "
            "Install one of: PyMuPDF (pip install pymupdf), "
            "pypdf (pip install pypdf), or pdfplumber (pip install pdfplumber)"
        )

    def _clean_pdf_text(self, text: str) -> str:
        """Clean extracted PDF text: normalize whitespace and remove artifacts."""
        lines = text.split("\n")
        cleaned = [line.strip() for line in lines if line.strip()]
        result = "\n".join(cleaned)
        # Collapse triple+ newlines to double
        result = re.sub(r"\n{3,}", "\n\n", result)
        return result

    async def _extract_docx_text(self, content: str) -> str:
        """Extract text from a DOCX document."""
        try:
            from io import BytesIO

            from docx import Document

            if self._looks_like_base64(content):
                docx_bytes = base64.b64decode(content)
            else:
                docx_bytes = content.encode()

            doc = Document(BytesIO(docx_bytes))
            text_parts: list[str] = []
            for para in doc.paragraphs:
                if para.text:
                    text_parts.append(para.text)

            return "\n\n".join(text_parts)
        except ImportError:
            raise

    def _looks_like_base64(self, content: str) -> bool:
        """Heuristic check for base64-encoded content."""
        if len(content) < 16:
            return False
        clean = content.replace(" ", "").replace("\n", "").replace("\r", "")
        if len(clean) % 4 != 0:
            return False
        return bool(re.match(r"^[A-Za-z0-9+/]*={0,2}$", clean))

    # ------------------------------------------------------------------
    # Contextual knowledge enrichment (CTX_KNOWLEDGE_ENABLED)
    # ------------------------------------------------------------------

    async def _enrich_chunk_with_context(
        self,
        chunk: str,
        full_document: str,
        content_type: str | None = None,
    ) -> str:
        """Enrich a chunk by generating surrounding context via LLM.

        Mirrors TypeScript ``getContextualizedChunks`` when CTX_KNOWLEDGE_ENABLED.
        """
        if not self._text_generation_provider:
            return chunk

        try:
            targets = self._get_context_targets(content_type, chunk)
            prompt = CONTEXTUAL_CHUNK_ENRICHMENT_PROMPT_TEMPLATE.format(
                doc_content=full_document[:8000],
                chunk_content=chunk,
                min_tokens=targets.min_tokens,
                max_tokens=targets.max_tokens,
            )

            enriched = await self._text_generation_provider.generate_text(prompt)
            if enriched and enriched.strip():
                return enriched.strip()
        except Exception as e:
            logger.warning(f"Context enrichment failed, using raw chunk: {e}")

        return chunk

    def _get_context_targets(
        self, content_type: str | None, chunk: str
    ) -> ContextTargets:
        """Select context targets based on content type and content heuristics."""
        if content_type:
            ct_lower = content_type.lower()
            if "pdf" in ct_lower:
                if self._contains_mathematical_content(chunk):
                    return CONTEXT_TARGETS["MATH_PDF"]
                return CONTEXT_TARGETS["PDF"]
            if any(
                lang in ct_lower
                for lang in ["javascript", "typescript", "python", "java", "code"]
            ):
                return CONTEXT_TARGETS["CODE"]
            if any(t in ct_lower for t in ["markdown", "html"]) or self._is_technical_doc(chunk):
                return CONTEXT_TARGETS["TECHNICAL"]

        return CONTEXT_TARGETS["DEFAULT"]

    def _contains_mathematical_content(self, text: str) -> bool:
        """Check if text contains mathematical notation."""
        math_patterns = [
            r"\$\$.+?\$\$",
            r"\$.+?\$",
            r"\\begin\{equation\}",
            r"\\sum_",
            r"\\int",
            r"\\frac\{",
        ]
        for pattern in math_patterns:
            if re.search(pattern, text, re.DOTALL):
                return True

        math_keywords = [
            "theorem", "lemma", "proof", "equation", "derivative",
            "integral", "matrix", "vector",
        ]
        lower = text.lower()
        count = sum(1 for kw in math_keywords if kw in lower)
        return count >= 2

    def _is_technical_doc(self, text: str) -> bool:
        """Check if text looks like technical documentation."""
        tech_patterns = [
            r"\b(?:version|v)\s*\d+\.\d+",
            r"\b(?:api|sdk|cli)\b",
            r"\b(?:GET|POST|PUT|DELETE)\b",
            r"\bREADME\b|\bCHANGELOG\b",
        ]
        for pattern in tech_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False

    # ------------------------------------------------------------------
    # Embeddings
    # ------------------------------------------------------------------

    async def _generate_fragment_embeddings(
        self,
        fragments: list[KnowledgeFragment],
    ) -> None:
        """Generate and attach embeddings to fragments."""
        if not self._embedding_provider:
            return

        texts = [f.content for f in fragments]
        results = await self._embedding_provider.generate_embeddings_batch(texts)

        for fragment, result in zip(fragments, results):
            fragment.embedding = result.embedding

    # ------------------------------------------------------------------
    # Search with cosine similarity
    # ------------------------------------------------------------------

    async def search(
        self,
        query: str,
        count: int = 10,
        threshold: float = 0.1,
    ) -> list[SearchResult]:
        """Search knowledge fragments using embedding-based cosine similarity.

        Falls back to keyword matching if no embedding provider is available.
        """
        if self._embedding_provider:
            return await self._search_with_embeddings(query, count, threshold)

        # Fallback: keyword-based search
        return self._search_with_keywords(query, count, threshold)

    async def _search_with_embeddings(
        self,
        query: str,
        count: int,
        threshold: float,
    ) -> list[SearchResult]:
        """Embedding-based similarity search."""
        result = await self._embedding_provider.generate_embedding(query)
        query_embedding = result.embedding

        results: list[SearchResult] = []

        for fragment in self._fragments.values():
            if fragment.embedding is None:
                continue

            similarity = self._cosine_similarity(query_embedding, fragment.embedding)

            if similarity >= threshold:
                document = self._documents.get(fragment.document_id)
                results.append(
                    SearchResult(
                        id=fragment.id,
                        content=fragment.content,
                        similarity=similarity,
                        document_id=fragment.document_id,
                        document_title=document.filename if document else None,
                        metadata=fragment.metadata,
                    )
                )

        results.sort(key=lambda x: x.similarity, reverse=True)
        return results[:count]

    def _search_with_keywords(
        self,
        query: str,
        count: int,
        threshold: float,
    ) -> list[SearchResult]:
        """Keyword-based fallback search."""
        query_lower = query.lower()
        query_words = query_lower.split()
        results: list[SearchResult] = []

        for fragment in self._fragments.values():
            content_lower = fragment.content.lower()
            matching = sum(1 for w in query_words if w in content_lower)

            if matching > 0:
                similarity = matching / len(query_words)
                if similarity >= threshold:
                    document = self._documents.get(fragment.document_id)
                    results.append(
                        SearchResult(
                            id=fragment.id,
                            content=fragment.content,
                            similarity=similarity,
                            document_id=fragment.document_id,
                            document_title=document.filename if document else None,
                            metadata=fragment.metadata,
                        )
                    )

        results.sort(key=lambda x: x.similarity, reverse=True)
        return results[:count]

    def _cosine_similarity(
        self,
        vec1: list[float],
        vec2: list[float],
    ) -> float:
        """Compute cosine similarity between two vectors."""
        if len(vec1) != len(vec2):
            return 0.0

        dot_product = sum(a * b for a, b in zip(vec1, vec2))
        norm1 = math.sqrt(sum(a * a for a in vec1))
        norm2 = math.sqrt(sum(b * b for b in vec2))

        if norm1 == 0 or norm2 == 0:
            return 0.0

        return dot_product / (norm1 * norm2)

    # ------------------------------------------------------------------
    # Get knowledge (convenience wrapper)
    # ------------------------------------------------------------------

    async def get_knowledge(
        self,
        query: str,
        count: int = 5,
    ) -> list[KnowledgeItem]:
        """Retrieve knowledge items for a query."""
        results = await self.search(query, count=count)

        return [
            KnowledgeItem(
                id=r.id,
                content=r.content,
                similarity=r.similarity,
                metadata=r.metadata,
            )
            for r in results
        ]

    # ------------------------------------------------------------------
    # RAG enrichment (mirrors TypeScript enrichConversationMemoryWithRAG)
    # ------------------------------------------------------------------

    async def enrich_conversation_memory_with_rag(
        self,
        memory_id: str,
        rag_metadata: RAGMetadata,
    ) -> None:
        """Enrich a conversation memory with RAG usage metadata."""
        if not self._memory_store:
            logger.debug("No memory store - skipping RAG enrichment")
            return

        try:
            existing = await self._memory_store.get_memory_by_id(memory_id)
            if not existing:
                logger.warning(f"Cannot enrich memory {memory_id} - not found")
                return

            updated_metadata = {
                **(existing.get("metadata", {})),
                "knowledgeUsed": True,
                "ragUsage": {
                    "retrievedFragments": [
                        {
                            "fragmentId": f.fragment_id,
                            "documentTitle": f.document_title,
                            "similarityScore": f.similarity_score,
                            "contentPreview": f.content_preview,
                        }
                        for f in rag_metadata.retrieved_fragments
                    ],
                    "queryText": rag_metadata.query_text,
                    "totalFragments": rag_metadata.total_fragments,
                    "retrievalTimestamp": rag_metadata.retrieval_timestamp,
                    "usedInResponse": rag_metadata.used_in_response,
                },
                "type": MemoryType.CUSTOM.value,
            }

            await self._memory_store.update_memory(
                {"id": memory_id, "metadata": updated_metadata}
            )
        except Exception as e:
            logger.warning(f"Failed to enrich memory {memory_id} with RAG data: {e}")

    def set_pending_rag_metadata(self, rag_metadata: RAGMetadata) -> None:
        """Queue RAG metadata for later enrichment of conversation memories."""
        now = int(time.time() * 1000)
        # Prune stale entries (older than 30 seconds)
        self._pending_rag_enrichment = [
            e for e in self._pending_rag_enrichment if now - e.timestamp < 30000
        ]
        self._pending_rag_enrichment.append(
            PendingRAGEntry(rag_metadata=rag_metadata, timestamp=now)
        )

    async def enrich_recent_memories_with_pending_rag(self) -> None:
        """Try to enrich recent conversation memories with queued RAG metadata."""
        if not self._pending_rag_enrichment or not self._memory_store:
            return

        try:
            recent_memories = await self._memory_store.get_memories(
                table_name="messages", count=10
            )
            now = int(time.time() * 1000)

            recent_convos = [
                m
                for m in recent_memories
                if m.get("metadata", {}).get("type") == "message"
                and now - (m.get("createdAt", 0) or 0) < 10000
                and not m.get("metadata", {}).get("ragUsage")
            ]
            recent_convos.sort(key=lambda m: m.get("createdAt", 0) or 0, reverse=True)

            to_remove: list[PendingRAGEntry] = []
            for entry in self._pending_rag_enrichment:
                for memory in recent_convos:
                    if (memory.get("createdAt", 0) or 0) > entry.timestamp:
                        mid = memory.get("id")
                        if mid:
                            await self.enrich_conversation_memory_with_rag(mid, entry.rag_metadata)
                            to_remove.append(entry)
                        break

            for entry in to_remove:
                if entry in self._pending_rag_enrichment:
                    self._pending_rag_enrichment.remove(entry)
        except Exception as e:
            logger.warning(f"Error enriching recent memories with RAG data: {e}")

    def build_rag_metadata(
        self,
        knowledge_items: list[KnowledgeItem],
        query_text: str,
    ) -> RAGMetadata:
        """Build RAG metadata from search results (used by the provider)."""
        return RAGMetadata(
            retrieved_fragments=[
                RetrievedFragmentInfo(
                    fragment_id=item.id,
                    document_title=str(item.metadata.get("filename", item.metadata.get("title", ""))),
                    similarity_score=item.similarity,
                    content_preview=item.content[:100] + "..." if len(item.content) > 100 else item.content,
                )
                for item in knowledge_items
            ],
            query_text=query_text,
            total_fragments=len(knowledge_items),
            retrieval_timestamp=int(time.time() * 1000),
        )

    # ------------------------------------------------------------------
    # CRUD helpers
    # ------------------------------------------------------------------

    async def check_existing_knowledge(self, knowledge_id: str) -> bool:
        """Check if a document already exists by ID."""
        return knowledge_id in self._documents

    async def delete_knowledge(self, document_id: str) -> bool:
        """Delete a document and all its fragments."""
        if document_id not in self._documents:
            return False

        document = self._documents[document_id]

        # Remove fragments
        for fragment in document.fragments:
            self._fragments.pop(fragment.id, None)

        # Remove content hash mapping
        if document.content_hash and document.content_hash in self._content_hashes:
            del self._content_hashes[document.content_hash]

        del self._documents[document_id]

        logger.info(f"Deleted document {document_id}")
        return True

    def get_documents(self) -> list[KnowledgeDocument]:
        """Return all stored documents."""
        return list(self._documents.values())

    def get_document(self, document_id: str) -> KnowledgeDocument | None:
        """Return a single document by ID."""
        return self._documents.get(document_id)

    def get_fragment_count(self) -> int:
        """Return total number of fragments across all documents."""
        return len(self._fragments)
